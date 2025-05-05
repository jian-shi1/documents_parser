import argparse
from borb.pdf import Document, Page, Paragraph, SingleColumnLayout, PDF
from constants import file_formats, archieve_formats, wordlist_filename
import docx
from numpy.random import choice, randint
import os
import pandas as pd
from pathlib import Path
import py7zr
from shutil import rmtree
import subprocess
import xlwt
import zipfile


class DocumentGenerator:
    # Лимиты на наполнение файлов
    _max_words_name = 3
    _max_paragraphs = 500
    _max_paragraph_words = 100
    _max_cell_words = 10
    _max_columns = 10
    _max_rows = 100
    # Лимиты на количество файлов
    _max_doc_num = 1
    _max_docx_num = 1
    _max_xls_num = 1
    _max_xlsx_num = 1
    _max_pdf_num = 1
    _max_zip_num = 1
    _max_7zip_num = 1
    _max_rar_num = 1

    banwords = set()
    word_bank = set()
    all_formats = set()

    def archieve_zip(workdir):
        sp = os.path.split(workdir)
        last_wd = sp[-1]
        with zipfile.ZipFile(f"{workdir}.zip", mode="w") as archive:
            for file_path in Path(workdir).iterdir():
                archive.write(file_path, arcname=file_path.name)

    def archieve_rar(workdir):
        sp = os.path.split(workdir)
        current_dir = Path(__file__).parent.absolute()
        last_wd = sp[-1]
        documents_dir = current_dir / sp[-2]
        dir_path = documents_dir / last_wd
        rar_path = documents_dir / f"{last_wd}.rar"
        try:
            subprocess.run([
                "rar",  # имя утилиты
                "a",  # архивируем
                "-ep1",  # не включаем внутри архива вложенности родительских директорий
                "-r",  # рекурсивно
                "-idq",  # без сообщений на вывод
                str(rar_path),  # название архива
                str(dir_path / "*")  # что архивируем
            ], check=True)
        except subprocess.CalledProcessError as e:
            print(f"RAR creation failed: {e}")
        except FileNotFoundError:
            print("'rar' command not found. Install RAR archiver first.")

    def archieve_7zip(workdir):
        sp = os.path.split(workdir)
        last_wd = sp[-1]
        with py7zr.SevenZipFile(f"{workdir}.7z", 'w') as archive:
            archive.writeall(workdir, arcname=os.path.basename(last_wd))

    do_archivation = {
        ".zip": archieve_zip,
        ".rar": archieve_rar,
        ".7z": archieve_7zip
    }

    def __init__(self):
        self.all_formats = file_formats | archieve_formats
        # банк слов из большого файла с английскими словами
        with open(wordlist_filename) as wb:
            self.word_bank = wb.read().splitlines()

    def _gen_wordlist(self, numwords):
        return choice(self.word_bank, numwords)

    def _gen_paragraph(self, is_cell=False):
        ic = int(is_cell)
        words_num = randint(1, ic * self._max_cell_words +
                            (1-ic) * self._max_paragraph_words)
        return ' '.join(self._gen_wordlist(words_num))

    def _gen_title(self, idx, ext, banwords):
        """
        Заголовок файла генерим с оглядкой на банворды
        После этого новый заголовок становится новым банвордом
        """
        num_words_in_title = randint(1, self._max_words_name + 1)
        s = '_'.join(self._gen_wordlist(num_words_in_title)) + str(idx) + ext
        while s in banwords:
            s = '_'.join(self._gen_wordlist(num_words_in_title)) + ext
        banwords.add(s)
        return s

    def _generate_doc(self, num, ext, banwords, workdir):
        """
        Для простоты doc-файл это набор параграфов случайных слов
        """
        for i in range(num):
            doc = docx.Document()
            name = os.path.join(workdir, self._gen_title(i, ext, banwords))
            num_pars = randint(1, self._max_paragraphs+1)
            for _ in range(num_pars):
                doc.add_paragraph(self._gen_paragraph())
            doc.save(name)

    def _generate_xls(self, num, banwords, workdir):
        """
        Для работы со старыми файлами вписываем слова в каждую ячейку по отдельности
        """
        ext = ".xls"
        for filenum in range(num):
            name = os.path.join(
                workdir, self._gen_title(filenum, ext, banwords))
            rows_num = randint(1, self._max_rows)
            cols_num = randint(1, self._max_columns)
            wb = xlwt.Workbook()
            ws = wb.add_sheet('Sheet 1')
            for i in range(rows_num):
                for j in range(cols_num):
                    ws.write(i, j, self._gen_paragraph(True))
            wb.save(name)

    def _generate_xlsx(self, num, banwords, workdir):
        """
        Генерация xlsx файла через генерацию датафрейма
        """
        ext = ".xlsx"
        for i in range(num):
            name = os.path.join(workdir, self._gen_title(i, ext, banwords))
            rows_num = randint(1, self._max_rows)
            cols_num = randint(1, self._max_columns)
            init_dict = {}
            for colname in self._gen_wordlist(cols_num):
                init_dict[colname] = self._gen_wordlist(rows_num)
            df = pd.DataFrame(init_dict)
            df.to_excel(name, index=False)

    def _generate_pdf(self, num, banwords, workdir):
        ext = ".pdf"
        for i in range(num):
            pdf = Document()
            page = Page()
            pdf.add_page(page)
            layout = SingleColumnLayout(page)

            name = os.path.join(workdir, self._gen_title(i, ext, banwords))
            num_pars = randint(1, self._max_paragraphs+1)
            for _ in range(num_pars):
                layout.add(Paragraph(self._gen_paragraph()))
            with open(name, "wb") as pdf_file_handle:
                PDF.dumps(pdf_file_handle, pdf)

    def _generate_archieve(self, num, ext, banwords, workdir, *, doc_n, docx_n, xls_n, xlsx_n, pdf_n):
        """
        Создадим папку, внутри которой сгенерим еще файлы
        Потом сожмем её в архив
        """
        for i in range(num):
            name = os.path.join(workdir, self._gen_title(i, '', banwords))
            Path(name).mkdir()
            self.generate(name, doc_num=docx_n, docx_num=docx_n, xls_num=xls_n,
                          xlsx_num=xlsx_n, pdf_num=pdf_n, szip_num=0, zip_num=0, rar_num=0)
            self.do_archivation[ext](name)
            rmtree(name)

    def generate(self, workdir, *, doc_num=_max_doc_num, docx_num=_max_docx_num, xls_num=_max_xls_num, xlsx_num=_max_xlsx_num,
                 pdf_num=_max_pdf_num, zip_num=_max_zip_num, rar_num=_max_rar_num, szip_num=_max_7zip_num):

        # банворды для нейминга это уже существующие в директории названия, будет пополняться
        existing_filenames = []
        for _, dirs, files in os.walk(workdir):
            existing_filenames.extend(files)
            existing_filenames.extend(dirs)
            break
        banwords = set([filename for filename in existing_filenames if os.path.splitext(
            filename)[1] in self.all_formats])

        existing_filenames = []
        for _, _, filenames in os.walk(workdir):
            existing_filenames.extend(filenames)
            break
        existing_filenames = set(existing_filenames)

        self._generate_doc(randint(1, doc_num+1), ".doc", banwords, workdir)
        self._generate_doc(randint(1, docx_num+1), ".docx", banwords, workdir)
        self._generate_xlsx(randint(1, xlsx_num+1), banwords, workdir)
        self._generate_xls(randint(1, xls_num+1), banwords, workdir)
        self._generate_pdf(randint(1, pdf_num+1), banwords, workdir)
        if rar_num > 0:
            self._generate_archieve(randint(1, szip_num+1),
                                    ".rar", banwords, workdir, doc_n=doc_num, docx_n=docx_num, xls_n=xls_num, xlsx_n=xlsx_num, pdf_n=pdf_num)

        if szip_num > 0:
            self._generate_archieve(randint(1, szip_num+1),
                                    ".7z", banwords, workdir, doc_n=doc_num, docx_n=docx_num, xls_n=xls_num, xlsx_n=xlsx_num, pdf_n=pdf_num)
        if zip_num > 0:
            self._generate_archieve(randint(1, szip_num+1),
                                    ".zip", banwords, workdir, doc_n=doc_num, docx_n=docx_num, xls_n=xls_num, xlsx_n=xlsx_num, pdf_n=pdf_num)


def main():
    arg_parser = argparse.ArgumentParser(
        description='Generate many files in dir')
    arg_parser.add_argument('dir', type=str, help='Input directory')
    args = arg_parser.parse_args()
    gen = DocumentGenerator()
    gen.generate(args.dir)


if __name__ == "__main__":
    main()
